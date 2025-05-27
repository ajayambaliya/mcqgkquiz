import os
import requests
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from telegram import Bot
import asyncio
import time
import re
from pymongo import MongoClient
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import parse_xml
from datetime import datetime

# Load credentials from environment variables
BOT_TOKEN = os.getenv('BOT_TOKEN')
TELEGRAM_CHANNEL_USERNAME = os.getenv('TELEGRAM_CHANNEL_USERNAME')
MONGO_CONNECTION_STRING = os.getenv('MONGO_CONNECTION_STRING')
DB_NAME = 'mcqgkdb'
COLLECTION_NAME = 'ScrapedLinks'

# Print credentials status (for debugging)
print(f"BOT_TOKEN: {'Set' if BOT_TOKEN else 'Not set'}")
print(f"TELEGRAM_CHANNEL_USERNAME: {TELEGRAM_CHANNEL_USERNAME}")
print(f"MONGO_CONNECTION_STRING: {'Set' if MONGO_CONNECTION_STRING else 'Not set'}")

# Connect to MongoDB
def get_mongo_client():
    client = MongoClient(MONGO_CONNECTION_STRING)
    return client[DB_NAME]

def get_stored_urls(collection):
    return set(doc['url'] for doc in collection[COLLECTION_NAME].find({"url": {"$exists": True}}))

def store_url(collection, url):
    collection[COLLECTION_NAME].update_one({'url': url}, {'$set': {'url': url}}, upsert=True)

def get_last_question_count(collection):
    config = collection['Config'].find_one({'key': 'last_question_count'})
    return config['value'] if config else 0

def store_last_question_count(collection, count):
    collection['Config'].update_one({'key': 'last_question_count'}, {'$set': {'key': 'last_question_count', 'value': count}}, upsert=True)

def should_reset_count():
    today = datetime.now()
    return today.day == 1  # Reset if it's the 1st of the month

# Function to fetch links from a URL
def fetch_links(url):
    try:
        response = requests.get(url)
        if response.status_code != 200:
            print(f"Failed to fetch URL {url}. Status code: {response.status_code}")
            return {}
        soup = BeautifulSoup(response.text, 'html.parser')
        # Find all divs with class 'home-post-item'
        post_items = soup.find_all('div', class_='home-post-item')
        if not post_items:
            print(f"Error: Could not find div with class 'home-post-item' on {url}")
            return {}
        # Extract links from <a> tags within <h3> elements
        links = {}
        for i, item in enumerate(post_items, 1):
            link_tag = item.find('h3').find('a', href=True) if item.find('h3') else None
            if link_tag:
                links[i] = link_tag['href']
        return links
    except Exception as e:
        print(f"Error fetching links from {url}: {e}")
        return {}

# Function to scrape content from the selected links
def scrape_content_from_links(selected_links):
    all_questions = []
    for link in selected_links:
        try:
            response = requests.get(link)
            if response.status_code != 200:
                print(f"Failed to fetch quiz page {link}. Status code: {response.status_code}")
                continue
            soup = BeautifulSoup(response.text, 'html.parser')
            # Find the quiz container
            quiz_container = soup.find('div', class_='ques_print')
            if quiz_container:
                questions = extract_questions(quiz_container)
                if questions:
                    all_questions.extend(questions)
                else:
                    print(f"No questions extracted from {link}.")
            else:
                print(f"No quiz content found on {link}. HTML snippet:")
                print(soup.prettify()[:1000])  # Print first 1000 chars for debugging
        except Exception as e:
            print(f"Error scraping {link}: {e}")
    return all_questions

# Function to extract questions, options, correct answers, and explanations
def extract_questions(quiz_container):
    questions = []
    # Find all question divs within the quiz container
    quiz_questions = quiz_container.find_all('div', class_='wp_quiz_question testclass')
    for quiz in quiz_questions:
        # Extract question text including the number
        question_text = quiz.text.strip()
        quesno_span = quiz.find('span', class_='quesno')
        if quesno_span:
            question_text = quesno_span.text.strip() + quiz.text.replace(quesno_span.text, '', 1).strip()

        # Extract options
        options_div = quiz.find_next('div', class_='wp_quiz_question_options')
        if not options_div:
            print(f"Warning: No options found for question: {question_text}")
            continue
        options_raw = options_div.get_text(separator='\n').split('\n')
        options = []
        for opt in options_raw:
            clean_option = re.sub(r'^\[.\]\s*', '', opt).strip()
            if clean_option:
                options.append(clean_option)

        # Extract correct answer
        answer_div = quiz.find_next('div', class_='wp_basic_quiz_answer')
        if not answer_div:
            print(f"Warning: No answer div found for question: {question_text}")
            continue
        correct_answer_div = answer_div.find('div', class_='ques_answer')
        if not correct_answer_div:
            print(f"Warning: No correct answer found for question: {question_text}")
            continue
        correct_answer_text = correct_answer_div.text.strip()
        # Extract the letter (e.g., 'D' from 'Correct Answer: D [Department of Science and Technology]')
        correct_answer_match = re.search(r'Correct Answer:\s*([A-D])', correct_answer_text)
        if not correct_answer_match:
            print(f"Warning: Could not parse correct answer for question: {question_text}")
            continue
        correct_answer_letter = correct_answer_match.group(1)
        try:
            correct_answer_index = ['A', 'B', 'C', 'D'].index(correct_answer_letter)
        except ValueError:
            print(f"Warning: Invalid correct answer letter '{correct_answer_letter}' for question: {question_text}")
            continue

        # Extract explanation
        explanation_div = answer_div.find('div', class_='answer_hint')
        explanation_text = explanation_div.text.replace('Notes:', '').strip() if explanation_div else "No explanation available."

        # Ensure valid question data
        if len(options) >= 2 and correct_answer_index != -1:
            questions.append({
                'question': question_text,
                'options': options,
                'correct_answer': correct_answer_index,
                'explanation': explanation_text
            })
        else:
            print(f"Skipping question due to invalid options or answer: {question_text}")

    return questions
def find_correct_answer_second_method(quiz):
    try:
        correct_answer_div = quiz.find('div', class_='correct_answer')
        correct_answer_letter = correct_answer_div.text.strip()[0]
        correct_answer_index = ['A', 'B', 'C', 'D'].index(correct_answer_letter)
        return correct_answer_index
    except:
        return -1

# Function to translate text to Gujarati
def translate_text(text, target_language='gu'):
    translator = GoogleTranslator(source='auto', target=target_language)
    translated_text = translator.translate(text)
    return translated_text

# Function to extract title from URL
def get_title_from_url(url):
    title = url.split('/')[-2] if url.endswith('/') else url.split('/')[-1]
    title = title.replace('-', ' ').title()
    return title

# Function to create a stylish Word file
def create_word_file(questions, start_count, last_url, filename="quiz_questions.docx"):
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Start with questions directly (no cover page)
    for idx, q in enumerate(questions, start=start_count + 1):
        question = translate_text(q['question'])
        options = [translate_text(opt) for opt in q['options']]
        correct_answer_idx = q['correct_answer']
        explanation = translate_text(q['explanation'])
        
        # Question text with number (e.g., "16. ...") in bold and colored
        q_text = doc.add_paragraph(question)
        q_text.runs[0].font.size = Pt(12)
        q_text.runs[0].font.color.rgb = RGBColor(75, 0, 130)  # Indigo
        q_text.runs[0].font.name = 'Calibri'
        q_text.runs[0].bold = True
        
        # Options without bullet points
        for i, option in enumerate(options):
            opt_para = doc.add_paragraph(f"{chr(65 + i)}. {option}")
            opt_para.runs[0].font.size = Pt(11)
            if i == correct_answer_idx:
                opt_para.runs[0].font.color.rgb = RGBColor(34, 139, 34)  # Forest Green
        
        # Answer in a shaded box
        ans_para = doc.add_paragraph()
        ans_run = ans_para.add_run(f"જવાબ: {options[correct_answer_idx]}")
        ans_run.font.size = Pt(11)
        ans_run.bold = True
        ans_run.font.color.rgb = RGBColor(255, 255, 255)  # White
        ans_para.paragraph_format.space_before = Pt(6)
        ans_para.paragraph_format.space_after = Pt(6)
        shading_elm = parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="228B22"/>')
        ans_para._element.get_or_add_pPr().append(shading_elm)
        
        # Explanation
        exp_para = doc.add_paragraph(f"સમજૂતી: {explanation}")
        exp_para.runs[0].font.size = Pt(10)
        exp_para.runs[0].italic = True
        exp_para.runs[0].font.color.rgb = RGBColor(105, 105, 105)  # Dim Gray
        
        # Decorative separator
        sep = doc.add_paragraph("✦✦✦")
        sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sep.runs[0].font.color.rgb = RGBColor(255, 105, 180)  # Hot Pink
    
    # Save the document with the title-based filename
    title = get_title_from_url(last_url)
    final_filename = f"{title}_{int(time.time())}.docx"
    doc.save(final_filename)
    return final_filename, idx

# Function to send the Word file to Telegram
async def send_word_file(filename, caption):
    bot = Bot(token=BOT_TOKEN)
    try:
        with open(filename, 'rb') as file:
            await bot.send_document(
                chat_id=TELEGRAM_CHANNEL_USERNAME,
                document=file,
                caption=caption
            )
        print(f"Sent {filename} to {TELEGRAM_CHANNEL_USERNAME}")
        os.remove(filename)
    except Exception as e:
        print(f"Error sending document: {e}")

# Main function
def main():
    url = "https://www.gktoday.in/gk-current-affairs-quiz-questions-answers/"
    links = fetch_links(url)

    db = get_mongo_client()
    stored_urls = get_stored_urls(db)

    # Check if today is the 1st of the month and reset count if so
    if should_reset_count():
        store_last_question_count(db, 0)
        print("Reset question count to 0 as it's the 1st of the month.")

    last_count = get_last_question_count(db)

    new_links = {num: link for num, link in links.items() if link not in stored_urls}
    
    if not new_links:
        print("No new links to scrape.")
        return

    all_questions = []
    last_url = None
    for link in new_links.values():
        print(f"Scraping link: {link}")
        questions = scrape_content_from_links([link])
        if questions:
            all_questions.extend(questions)
            store_url(db, link)
            last_url = link
        else:
            print(f"No questions found on {link}.")

    if all_questions:
        title = get_title_from_url(last_url)
        filename = f"{title}_{int(time.time())}.docx"
        caption = f"{title} ગુજરાતીમાં"
        
        filename, new_last_count = create_word_file(all_questions, last_count, last_url, filename)
        asyncio.run(send_word_file(filename, caption))
        store_last_question_count(db, new_last_count)
    else:
        print("No questions to process.")

if __name__ == "__main__":
    main()
